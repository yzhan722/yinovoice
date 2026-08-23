"""Export the Stage1 capability report to standalone HTML and DOCX."""

from __future__ import annotations

import html
import re
from pathlib import Path

import markdown
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SRC = Path(__file__).with_name("2026-08-18-stage1-capability-report.md")
HTML_OUT = SRC.with_suffix(".html")
DOCX_OUT = SRC.with_suffix(".docx")

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x1D, 0x4E, 0x89)
MUTED = RGBColor(0x4B, 0x55, 0x63)
HEADER_FILL = "1F3A5F"
ALT_FILL = "F4F7FB"
GREEN = RGBColor(0x1B, 0x7A, 0x4A)
AMBER = RGBColor(0x9A, 0x67, 0x09)
RED = RGBColor(0xB4, 0x23, 0x18)


def md_to_body_html(text: str) -> str:
    return markdown.markdown(
        text,
        extensions=["tables", "fenced_code", "nl2br", "sane_lists"],
    )


def wrap_html(title: str, body: str) -> str:
    return f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1" />
  <title>{html.escape(title)}</title>
  <style>
    :root {{
      --ink: #1a2332;
      --muted: #4b5563;
      --line: #d7dee8;
      --navy: #1f3a5f;
      --paper: #f7f5f0;
      --card: #ffffff;
      --ok: #1b7a4a;
      --warn: #9a6709;
      --no: #b42318;
    }}
    * {{ box-sizing: border-box; }}
    body {{
      margin: 0;
      background: var(--paper);
      color: var(--ink);
      font: 16px/1.7 "Source Han Sans SC", "Noto Sans SC", "Microsoft YaHei",
        "PingFang SC", sans-serif;
    }}
    .page {{
      max-width: 880px;
      margin: 32px auto 64px;
      padding: 40px 48px 56px;
      background: var(--card);
      border: 1px solid var(--line);
    }}
    header.cover {{
      border-bottom: 3px solid var(--navy);
      padding-bottom: 20px;
      margin-bottom: 28px;
    }}
    .kicker {{
      letter-spacing: 0.12em;
      font-size: 12px;
      color: var(--navy);
      font-weight: 700;
    }}
    h1 {{
      font-size: 28px;
      line-height: 1.3;
      margin: 8px 0 10px;
      color: var(--navy);
    }}
    h2 {{
      font-size: 20px;
      color: var(--navy);
      margin: 36px 0 12px;
      padding-top: 8px;
      border-top: 1px solid var(--line);
    }}
    h3 {{
      font-size: 16px;
      margin: 24px 0 8px;
      color: #24344d;
    }}
    p, li {{ color: var(--ink); }}
    ul, ol {{ padding-left: 1.3em; }}
    blockquote {{
      margin: 16px 0;
      padding: 10px 16px;
      border-left: 4px solid var(--navy);
      background: #eef3f8;
      color: #24344d;
    }}
    code {{
      font-family: Consolas, "Sarasa Mono SC", monospace;
      font-size: 0.9em;
      background: #eef2f6;
      padding: 0 4px;
    }}
    pre {{
      background: #1a2332;
      color: #e8edf4;
      padding: 14px 16px;
      overflow: auto;
      font-size: 13px;
      line-height: 1.55;
    }}
    pre code {{ background: none; color: inherit; padding: 0; }}
    table {{
      width: 100%;
      border-collapse: collapse;
      font-size: 14px;
      margin: 12px 0 20px;
    }}
    th, td {{
      border: 1px solid var(--line);
      padding: 8px 10px;
      vertical-align: top;
      text-align: left;
    }}
    th {{
      background: var(--navy);
      color: #fff;
      font-weight: 600;
    }}
    tr:nth-child(even) td {{ background: #f4f7fb; }}
    hr {{ border: 0; border-top: 1px solid var(--line); margin: 28px 0; }}
    em.note, .footer {{
      color: var(--muted);
      font-size: 13px;
    }}
    @media print {{
      body {{ background: #fff; }}
      .page {{
        margin: 0;
        border: 0;
        max-width: none;
        padding: 16mm 16mm 18mm;
      }}
    }}
  </style>
</head>
<body>
  <article class="page">
    <header class="cover">
      <div class="kicker">YINO VOICE · STAGE1 · 内部汇报</div>
    </header>
    {body}
  </article>
</body>
</html>
"""


INLINE_RE = re.compile(
    r"`([^`]+)`|\*\*(.+?)\*\*|_(.+?)_|([^`*_]+)",
)


def set_run_font(run, *, east_asia="微软雅黑", ascii_font="Calibri", size=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    if color is not None:
        run.font.color.rgb = color


def add_runs(paragraph, text: str, *, size=11, color=MUTED):
    for match in INLINE_RE.finditer(text):
        code, bold, italic, plain = match.groups()
        if code is not None:
            run = paragraph.add_run(code)
            set_run_font(run, ascii_font="Consolas", east_asia="微软雅黑", size=size - 1, color=ACCENT)
        elif bold is not None:
            run = paragraph.add_run(bold)
            set_run_font(run, size=size, bold=True, color=NAVY)
        elif italic is not None:
            run = paragraph.add_run(italic)
            set_run_font(run, size=size, color=color)
            run.italic = True
        else:
            run = paragraph.add_run(plain)
            set_run_font(run, size=size, color=color)


def shade_cell(cell, fill: str):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): fill,
        },
    )
    tc_pr.append(shd)


def set_cell_text(cell, text: str, *, header=False, size=10):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    add_runs(p, text, size=size, color=RGBColor(0xFF, 0xFF, 0xFF) if header else MUTED)
    if header:
        for run in p.runs:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def status_color(text: str) -> RGBColor | None:
    if any(key in text for key in ("已交付", "已做", "真实")):
        return GREEN
    if any(key in text for key in ("未做", "不能", "不上")):
        return RED
    if any(key in text for key in ("部分", "演示", "半交付", "未全部", "Prompt 约束")):
        return AMBER
    return None


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            header = r_idx == 0
            set_cell_text(cell, value, header=header, size=10)
            shade_cell(cell, HEADER_FILL if header else (ALT_FILL if r_idx % 2 == 0 else "FFFFFF"))
            if not header:
                tone = status_color(value)
                if tone is not None:
                    for run in cell.paragraphs[0].runs:
                        run.font.color.rgb = tone
                        run.bold = True
    doc.add_paragraph()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].startswith("|"):
        raw = lines[i].strip()
        cells = [c.strip() for c in raw.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def build_docx(md_text: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    header = section.header.paragraphs[0]
    header_run = header.add_run("Yino Voice Stage1 能力范围汇报  ·  内部材料  ·  2026-08-18")
    set_run_font(header_run, size=9, color=ACCENT)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("不上生产  ·  仅描述已落地能力与边界")
    set_run_font(footer_run, size=9, color=MUTED)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run("\n".join(code_buf))
                set_run_font(run, ascii_font="Consolas", east_asia="微软雅黑", size=9, color=NAVY)
                p.paragraph_format.left_indent = Cm(0.4)
                code_buf = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and lines[i + 1].startswith("|"):
            rows, i = parse_table(lines, i)
            if rows:
                add_table(doc, rows)
            continue

        if line.strip() == "---":
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=0)
            for run in p.runs:
                set_run_font(run, size=22, bold=True, color=NAVY)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:].strip(), level=1)
            for run in p.runs:
                set_run_font(run, size=16, bold=True, color=NAVY)
            i += 1
            continue
        if line.startswith("### "):
            p = doc.add_heading(line[4:].strip(), level=2)
            for run in p.runs:
                set_run_font(run, size=13, bold=True, color=ACCENT)
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            add_runs(p, line[2:].strip(), size=12, color=NAVY)
            for run in p.runs:
                run.italic = True
            i += 1
            continue

        m = re.match(r"^(\d+)\.\s+(.*)$", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            add_runs(p, m.group(2), size=11, color=MUTED)
            i += 1
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:], size=11, color=MUTED)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, line.strip(), size=11, color=MUTED)
        i += 1

    return doc


def main() -> None:
    md_text = SRC.read_text(encoding="utf-8")
    title = "Yino Voice Stage1 能力范围汇报"
    body = md_to_body_html(md_text)
    # Drop the duplicated H1 inside body cover; keep document h1.
    HTML_OUT.write_text(wrap_html(title, body), encoding="utf-8")
    build_docx(md_text).save(DOCX_OUT)
    print(HTML_OUT)
    print(DOCX_OUT)


if __name__ == "__main__":
    main()
